"""
import docx
from loaders.file_loader import FileLoader

class DOCXLoader(FileLoader):
    def validate_file(self):
        if not self.file_path.endswith('.docx'):
            raise ValueError("Invalid file type. Expected a DOCX file.")

    def load_file(self):
        return docx.Document(self.file_path)
        """

import docx
import io
import msoffcrypto
from loaders.file_loader import FileLoader
from exceptions.content_access_exceptions import ContentAccessError

class DOCXLoader(FileLoader):
    def __init__(self, file_path, password=None):
        super().__init__(file_path, password)

    def validate_file(self):
        if not self.file_path.endswith('.docx'):
            raise ValueError("Invalid file type. Expected a DOCX file.")

    def load_file(self):
        # return docx.Document(self.file_path)
        # Attempt to open the DOCX file, handling encrypted files if necessary
        try:
            # Open the DOCX file as a binary file
            with open(self.file_path, "rb") as file:
                office_file = msoffcrypto.OfficeFile(file)

                # Check if the file is encrypted
                if office_file.is_encrypted():
                    if self.password is None:
                        raise ValueError("The DOCX file is password-protected. A password is required to open it.")

                    # Load the key with the provided password
                    office_file.load_key(password=self.password)

                    # Decrypt to an in-memory file
                    decrypted_file = io.BytesIO()
                    office_file.decrypt(decrypted_file)
                    decrypted_file.seek(0)  # Reset the pointer to the start of the file

                    # Load the decrypted DOCX content
                    return docx.Document(decrypted_file)

                else:
                    # If the file is not encrypted, load normally
                    return docx.Document(file)
        except Exception as e:
            raise ValueError(f"Failed to load DOCX file: {e}")
        
    def verify_content(self, Document):
        return super().verify_content(Document)

    def get_metadata(self):
        # Load the document
        doc = self.load_file()
        
        # Initialize a dictionary to store metadata
        metadata = {
            "title": doc.core_properties.title,
            "author": doc.core_properties.author,
            "subject": doc.core_properties.subject,
            "keywords": doc.core_properties.keywords,
            "comments": doc.core_properties.comments,
            "last_modified_by": doc.core_properties.last_modified_by,
            "created": doc.core_properties.created,
            "modified": doc.core_properties.modified,
            "category": doc.core_properties.category,
            "content_status": doc.core_properties.content_status,
            "identifier": doc.core_properties.identifier,
            "language": doc.core_properties.language,
            "version": doc.core_properties.version
        }

        # Filter out None values
        metadata = {k: v for k, v in metadata.items() if v is not None}

        return metadata
    
    def get_fileType(self):
        return super().get_fileType()