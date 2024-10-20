import pytest
from loaders.docx_loader import DOCXLoader
from pathlib import Path

@pytest.fixture
def sample_docx_file(tmp_path):
    # Create a sample DOCX file for testing purposes
    sample_file = tmp_path / "sample.docx"
    from docx import Document
    doc = Document()
    doc.add_paragraph("This is a sample paragraph.")
    doc.add_paragraph("Another paragraph with a hyperlink.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Cell 1"
    table.cell(0, 1).text = "Cell 2"
    table.cell(1, 0).text = "Cell 3"
    table.cell(1, 1).text = "Cell 4"
    doc.save(sample_file)
    return sample_file

class TestDOCXLoader:
    
    def test_validate_file(self, sample_docx_file):
        """Test that the DOCX file is validated correctly"""
        loader = DOCXLoader(sample_docx_file)
        try:
            loader.validate_file()
        except Exception:
            pytest.fail("validate_file() raised an unexpected exception.")

    def test_invalid_file_extension(self, tmp_path):
        """Test that an invalid file extension raises a ValueError"""
        invalid_file = tmp_path / "invalid.txt"
        invalid_file.write_text("This is not a DOCX file.")
        
        with pytest.raises(ValueError, match="is not a valid DOCX file"):
            loader = DOCXLoader(invalid_file)
            loader.validate_file()

    def test_load_file(self, sample_docx_file):
        """Test that the DOCX file is loaded correctly"""
        loader = DOCXLoader(sample_docx_file)
        loader.validate_file()
        loader.load_file()
        
        assert loader.document is not None

    def test_get_text(self, sample_docx_file):
        """Test that text is extracted correctly from the DOCX file"""
        loader = DOCXLoader(sample_docx_file)
        loader.validate_file()
        loader.load_file()
        
        text = loader.get_text()
        assert "This is a sample paragraph." in text
        assert "Another paragraph with a hyperlink." in text

    def test_get_tables(self, sample_docx_file):
        """Test that tables are extracted correctly from the DOCX file"""
        loader = DOCXLoader(sample_docx_file)
        loader.validate_file()
        loader.load_file()

        tables = loader.get_tables()
        assert len(tables) == 1  # There is one table in the document
        assert tables[0] == [["Cell 1", "Cell 2"], ["Cell 3", "Cell 4"]]

    def test_get_links(self, sample_docx_file):
        """Test that links are extracted correctly from the DOCX file"""
        loader = DOCXLoader(sample_docx_file)
        loader.validate_file()
        loader.load_file()

        links = loader.get_links()
        assert links == []  # Since the test document doesn't actually contain links

    def test_get_images(self, sample_docx_file):
        """Test that image extraction works (placeholder for now)"""
        loader = DOCXLoader(sample_docx_file)
        loader.validate_file()
        loader.load_file()

        images = loader.get_images()
        assert images == "Image extraction not supported with the current library. Use alternative methods."